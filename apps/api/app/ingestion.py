from __future__ import annotations

import io
import json
import os
import mimetypes
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import PurePosixPath
from typing import Any


MAX_FILE_BYTES = 512 * 1024 * 1024
MAX_ARCHIVE_FILES = 500
MAX_ARCHIVE_BYTES = 2 * 1024 * 1024 * 1024
# Render's Starter instance has a 512 MB memory ceiling. Do not materialize
# very large drawing members or PDFs in Python just to extract a small preview.
MAX_IN_MEMORY_MEMBER_BYTES = 8 * 1024 * 1024
TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".ifc", ".dxf", ".html", ".log"}
DRAWING_EXTENSIONS = {".pdf", ".dwg", ".dxf", ".rvt", ".rfa", ".ifc", ".png", ".jpg", ".jpeg", ".tif", ".tiff"}


def _drawing_hints(name: str, text: str) -> dict[str, Any]:
    haystack = f"{name} {text}".lower()
    disciplines: list[str] = []
    keywords = {
        "architectural": ("architectural", "floor plan", "elevation", "section", "a-", "arch"),
        "mep": ("hvac", "mechanical", "plumbing", "mep", "m-", "duct", "chiller"),
        "electrical": ("electrical", "lighting", "power", "e-", "photovoltaic", "solar"),
        "structural": ("structural", "foundation", "column", "beam", "s-"),
        "landscape": ("landscape", "planting", "irrigation", "native", "habitat", "l-"),
        "civil": ("civil", "stormwater", "grading", "site plan", "c-"),
    }
    for discipline, terms in keywords.items():
        if any(term in haystack for term in terms):
            disciplines.append(discipline)
    sheets = sorted(set(re.findall(r"\b(?:[A-Z]{1,3}-?\d{1,3}(?:\.\d+)?)\b", name.upper() + " " + text.upper())))[:30]
    return {"is_drawing_candidate": name.lower().endswith(tuple(DRAWING_EXTENSIONS)), "disciplines": disciplines, "sheet_labels": sheets, "keyword_hits": [k for k in ("HVAC", "fresh air", "energy model", "EPD", "daylight", "biodiversity", "embodied carbon", "resilience", "EV charging") if k.lower() in haystack]}


def _extract_pdf(data: bytes) -> tuple[str, int, list[str]]:
    warnings: list[str] = []
    if len(data) > MAX_IN_MEMORY_MEMBER_BYTES:
        return "", 0, ["Large PDF indexed in metadata mode; text preview is deferred to keep online processing stable."]
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data), strict=False)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                warnings.append("PDF is encrypted; upload an unlocked copy for text extraction.")
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        return text, len(reader.pages), warnings
    except Exception as exc:
        warnings.append(f"PDF text extraction unavailable ({type(exc).__name__}); page/image metadata retained.")
        return "", 0, warnings


def _extract_pdf_bounded_path(path: str) -> tuple[str, int, list[str]]:
    """Extract a small PDF preview in a child process so malformed PDFs cannot exhaust the API worker."""
    script = (
        "import json,sys\n"
        "try:\n"
        " import resource; resource.setrlimit(resource.RLIMIT_AS,(256*1024*1024,256*1024*1024))\n"
        "except Exception: pass\n"
        "try:\n"
        " from pypdf import PdfReader\n"
        " with open(sys.argv[1],'rb') as f:\n"
        "  r=PdfReader(f,strict=False); pages=r.pages; n=min(3,len(pages)); text='\\n'.join((pages[i].extract_text() or '')[:20000] for i in range(n))\n"
        " print(json.dumps({'text':text[:100000],'pages':len(pages)}))\n"
        "except Exception as exc: print(json.dumps({'text':'','pages':0,'error':type(exc).__name__}))\n"
    )
    try:
        result = subprocess.run([sys.executable, "-c", script, path], capture_output=True, text=True, timeout=12, check=False)
        payload = json.loads(result.stdout.strip() or "{}")
        warnings = []
        if payload.get("pages", 0) > 3:
            warnings.append("Text preview extracted from the first 3 pages; full PDF page count retained.")
        if payload.get("error"):
            warnings.append("PDF preview was limited to protect online processing; filename and drawing metadata were retained.")
        return str(payload.get("text", "")), int(payload.get("pages", 0)), warnings
    except Exception as exc:
        return "", 0, [f"PDF preview unavailable ({type(exc).__name__}); filename and drawing metadata were retained."]


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        document = Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows[:30]:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)[:100_000]
    except Exception:
        return ""


def _extract_spreadsheet(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(min_row=1, max_row=30, values_only=True):
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if values:
                    parts.append(" | ".join(values[:20]))
        workbook.close()
        return "\n".join(parts)[:100_000]
    except Exception:
        return ""


def extract_file(name: str, data: bytes, content_type: str | None = None, archive_member: str | None = None, *, declared_size: int | None = None, skip_content: bool = False, text_override: str | None = None, page_count_override: int | None = None, warnings_override: list[str] | None = None) -> dict[str, Any]:
    if len(data) > MAX_FILE_BYTES:
        raise ValueError(f"{name} exceeds the {MAX_FILE_BYTES // (1024 * 1024)} MB file limit")
    clean_name = PurePosixPath(name.replace("\\", "/")).name or "upload.bin"
    ext = PurePosixPath(clean_name).suffix.lower()
    text, page_count, warnings = "", 0, []
    if text_override is not None:
        text, page_count, warnings = text_override, page_count_override or 0, list(warnings_override or [])
    elif skip_content:
        # Large archive members are indexed from their filename and ZIP
        # metadata without adding a user-facing error to the document card.
        pass
    elif ext in TEXT_EXTENSIONS:
        text = data[:2_000_000].decode("utf-8", errors="replace")
    elif ext == ".pdf":
        text, page_count, warnings = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    elif ext in {".xlsx", ".xlsm"}:
        text = _extract_spreadsheet(data)
    elif ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        warnings.append("Image received; OCR/vision provider is not configured, so recognition uses filename metadata.")
    elif ext in {".dwg", ".rvt", ".rfa"}:
        warnings.append(f"{ext.upper()} geometry is retained as metadata; connect a CAD/BIM parser for geometry-level quantities.")
    hints = _drawing_hints(clean_name, text)
    return {"filename": clean_name, "archive_member": archive_member, "mime_type": content_type or mimetypes.guess_type(clean_name)[0] or "application/octet-stream", "extension": ext, "size_bytes": declared_size if declared_size is not None else len(data), "text": text[:100_000], "page_count": page_count, "warnings": warnings, "drawing": hints}


def extract_upload(name: str, data: bytes, content_type: str | None = None) -> list[dict[str, Any]]:
    if len(data) > MAX_FILE_BYTES:
        raise ValueError(f"{name} exceeds the {MAX_FILE_BYTES // (1024 * 1024)} MB file limit")
    if not name.lower().endswith(".zip"):
        return [extract_file(name, data, content_type)]
    results: list[dict[str, Any]] = []
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        members = [m for m in archive.infolist() if not m.is_dir()]
        if len(members) > MAX_ARCHIVE_FILES:
            raise ValueError(f"ZIP contains more than {MAX_ARCHIVE_FILES} files")
        total = 0
        for member in members:
            path = PurePosixPath(member.filename.replace("\\", "/"))
            if path.is_absolute() or ".." in path.parts:
                raise ValueError(f"Unsafe ZIP path rejected: {member.filename}")
            if member.file_size > MAX_FILE_BYTES or total + member.file_size > MAX_ARCHIVE_BYTES:
                raise ValueError("ZIP uncompressed size exceeds the safe processing limit")
            total += member.file_size
            if path.suffix.lower() == ".pdf" and member.file_size > MAX_IN_MEMORY_MEMBER_BYTES:
                with tempfile.NamedTemporaryFile(suffix=".pdf") as temp:
                    with archive.open(member, "r") as source:
                        while chunk := source.read(1024 * 1024):
                            temp.write(chunk)
                    temp.flush()
                    text, pages, warnings = _extract_pdf_bounded_path(temp.name)
                results.append(extract_file(str(path), b"", None, archive_member=str(path), declared_size=member.file_size, text_override=text, page_count_override=pages, warnings_override=warnings))
            elif member.file_size > MAX_IN_MEMORY_MEMBER_BYTES:
                results.append(extract_file(str(path), b"", None, archive_member=str(path), declared_size=member.file_size, skip_content=True))
            else:
                payload = archive.read(member)
                results.append(extract_file(str(path), payload, None, archive_member=str(path)))
    return results


def extract_upload_path(name: str, path: str, content_type: str | None = None) -> list[dict[str, Any]]:
    """Extract a ZIP incrementally from disk so large archives do not duplicate in RAM."""
    if not name.lower().endswith(".zip"):
        with open(path, "rb") as handle:
            return [extract_file(name, handle.read(), content_type)]
    results: list[dict[str, Any]] = []
    with zipfile.ZipFile(path) as archive:
        members = [m for m in archive.infolist() if not m.is_dir()]
        if len(members) > MAX_ARCHIVE_FILES:
            raise ValueError(f"ZIP contains more than {MAX_ARCHIVE_FILES} files")
        total = 0
        for member in members:
            member_path = PurePosixPath(member.filename.replace("\\", "/"))
            if member_path.is_absolute() or ".." in member_path.parts:
                raise ValueError(f"Unsafe ZIP path rejected: {member.filename}")
            if member.file_size > MAX_FILE_BYTES or total + member.file_size > MAX_ARCHIVE_BYTES:
                raise ValueError("ZIP uncompressed size exceeds the safe processing limit")
            total += member.file_size
            if member_path.suffix.lower() == ".pdf" and member.file_size > MAX_IN_MEMORY_MEMBER_BYTES:
                with tempfile.NamedTemporaryFile(suffix=".pdf") as temp:
                    with archive.open(member, "r") as source:
                        while chunk := source.read(1024 * 1024):
                            temp.write(chunk)
                    temp.flush()
                    text, pages, warnings = _extract_pdf_bounded_path(temp.name)
                results.append(extract_file(str(member_path), b"", None, archive_member=str(member_path), declared_size=member.file_size, text_override=text, page_count_override=pages, warnings_override=warnings))
            elif member.file_size > MAX_IN_MEMORY_MEMBER_BYTES:
                results.append(extract_file(str(member_path), b"", None, archive_member=str(member_path), declared_size=member.file_size, skip_content=True))
            else:
                with archive.open(member, "r") as handle:
                    results.append(extract_file(str(member_path), handle.read(), None, archive_member=str(member_path)))
    return results
